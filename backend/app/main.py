from contextlib import asynccontextmanager
from io import BytesIO

from docx import Document
from fastapi import BackgroundTasks, FastAPI, File, Form, Header, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from fastapi.responses import StreamingResponse
from sqlalchemy import inspect, text

from app.api.auth import router as auth_router
from app.api.files import router as files_router
from app.api.objects import router as objects_router
from app.core.config import get_settings
from app.core.logging import (
    configure_logging,
    elapsed_ms,
    get_logger,
    get_request_identity,
    log_audit_event,
    start_timer,
)
from app.core.telemetry import record_request
from app.db.base import Base
from app.db.session import SessionLocal, engine
from app.models import files, objects, user  # noqa: F401
from app.schemas.pipeline import (
    PipelineConfig,
    PipelinePreprocessRefreshRequest,
    PipelinePreprocessRefreshResponse,
    PipelineProfileStoredResponse,
    PipelineRequest,
    PipelineRunResponse,
    PipelineStoredProfileRequest,
    PipelineStoredRunRequest,
    ReportRequest,
)
from app.services.pipeline_engine import (
    fetch_stored_dataset_profile,
    fetch_system_dashboard,
    refresh_preprocessing_from_storage,
    run_pipeline_from_storage,
    run_pipeline_via_services,
    upload_and_profile_dataset,
)
from app.services.user_service import bootstrap_admin

settings = get_settings()
configure_logging()
request_logger = get_logger("app.request")


def initialize_application_state() -> None:
    Base.metadata.create_all(bind=engine)
    inspector = inspect(engine)
    if "comparison_history" in inspector.get_table_names():
        existing = {column["name"] for column in inspector.get_columns("comparison_history")}
        migrations = {
            "project_id": "ALTER TABLE comparison_history ADD COLUMN project_id INTEGER",
            "version_number": "ALTER TABLE comparison_history ADD COLUMN version_number INTEGER DEFAULT 1",
            "parent_history_id": "ALTER TABLE comparison_history ADD COLUMN parent_history_id INTEGER",
            "tags_json": "ALTER TABLE comparison_history ADD COLUMN tags_json TEXT",
        }
        with engine.begin() as connection:
            for column_name, ddl in migrations.items():
                if column_name not in existing:
                    connection.execute(text(ddl))

    db = SessionLocal()
    try:
        bootstrap_admin(db)
    finally:
        db.close()


@asynccontextmanager
async def lifespan(_: FastAPI):
    initialize_application_state()
    yield


app = FastAPI(
    title=settings.app_name,
    version="0.2.0",
    debug=settings.debug,
    description="Модульный монолит для сравнительного анализа объектов.",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_origin_list,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def request_logging_middleware(request: Request, call_next):
    started_at = start_timer()
    identity = get_request_identity(request)
    request_id = request.headers.get("x-request-id") or f"req-{int(started_at * 1000)}"
    request.state.request_id = request_id
    request.state.user_id = identity.get("user_id")
    request.state.user_email = identity.get("user_email")
    request.state.user_role = identity.get("user_role")

    try:
        response = await call_next(request)
    except Exception:
        duration_ms = elapsed_ms(started_at)
        request_logger.exception(
            "request_failed %s",
            " ".join(
                [
                    f"request_id={request_id!r}",
                    f"method={request.method!r}",
                    f"path={request.url.path!r}",
                    f"query={str(request.url.query)!r}",
                    f"user_id={identity.get('user_id')!r}",
                    f"user_email={identity.get('user_email')!r}",
                    f"duration_ms={duration_ms:.2f}",
                ]
            ),
        )
        raise

    duration_ms = elapsed_ms(started_at)
    response.headers["X-Request-Id"] = request_id
    request_logger.info(
        "request_completed request_id=%r method=%r path=%r query=%r status_code=%r user_id=%r user_email=%r duration_ms=%.2f",
        request_id,
        request.method,
        request.url.path,
        str(request.url.query),
        response.status_code,
        identity.get("user_id"),
        identity.get("user_email"),
        duration_ms,
    )
    record_request(request.url.path, duration_ms, response.status_code)
    return response


@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException) -> JSONResponse:
    if exc.status_code >= 500:
        request_logger.error(
            "http_exception request_id=%r method=%r path=%r status_code=%r detail=%r",
            getattr(request.state, "request_id", None),
            request.method,
            request.url.path,
            exc.status_code,
            exc.detail,
            exc_info=True,
        )
    else:
        request_logger.warning(
            "http_exception request_id=%r method=%r path=%r status_code=%r detail=%r",
            getattr(request.state, "request_id", None),
            request.method,
            request.url.path,
            exc.status_code,
            exc.detail,
        )
    return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})


@app.exception_handler(Exception)
async def unhandled_exception_handler(request: Request, exc: Exception) -> JSONResponse:
    request_logger.error(
        "unhandled_exception request_id=%r method=%r path=%r user_id=%r user_email=%r",
        getattr(request.state, "request_id", None),
        request.method,
        request.url.path,
        getattr(request.state, "user_id", None),
        getattr(request.state, "user_email", None),
        exc_info=True,
    )
    return JSONResponse(
        status_code=500,
        content={"detail": {"code": "INTERNAL_SERVER_ERROR", "message": "Внутренняя ошибка сервера"}},
    )


@app.get("/health")
def healthcheck() -> dict[str, str]:
    return {"status": "ok"}


SUMMARY_LABELS = {
    "objects_count": "Количество объектов",
    "criteria_count": "Количество критериев",
    "weights_sum": "Сумма весов",
    "best_object_id": "Лучший объект",
    "best_score": "Лучшая оценка",
    "normalization_notes": "Нормализация",
    "mode": "Режим анализа",
    "target_object_id": "Целевой объект",
    "confidence_score": "Надежность расчета",
    "confidence_notes": "Замечания к качеству",
    "ranking_stability_note": "Устойчивость рейтинга",
    "analog_groups": "Группы аналогов",
    "dominance_pairs": "Доминирование объектов",
}

DIRECTION_LABELS = {
    "maximize": "Максимизация",
    "minimize": "Минимизация",
    "target": "Близость к целевому значению",
}

MODE_LABELS = {
    "rating": "Рейтинг объектов",
    "analog_search": "Поиск аналогов",
}


def report_label(key: str) -> str:
    return SUMMARY_LABELS.get(key, key.replace("_", " ").title())


def report_value(key: str, value: object) -> str:
    if value is None:
        return "-"
    if key == "mode":
        return MODE_LABELS.get(str(value), str(value))
    if isinstance(value, list):
        if not value:
            return "Нет данных"
        if all(isinstance(item, str) for item in value):
            return "; ".join(str(item) for item in value)
        return f"{len(value)} записей"
    if isinstance(value, float):
        return f"{value:.4f}"
    if isinstance(value, dict):
        return "См. детальный раздел"
    return str(value)


@app.get(f"{settings.api_prefix}/system/dashboard")
async def system_dashboard(authorization: str | None = Header(default=None)) -> dict:
    return await fetch_system_dashboard(authorization=authorization)


@app.post(f"{settings.api_prefix}/pipeline/upload-profile", response_model=PipelineProfileStoredResponse)
async def pipeline_upload_profile(
    file: UploadFile = File(...),
    detail_level: str = Form(default="summary"),
) -> PipelineProfileStoredResponse:
    try:
        body = await file.read()
        data = await upload_and_profile_dataset(
            filename=file.filename or "dataset",
            body=body,
            detail_level="detailed" if detail_level == "detailed" else "summary",
        )
        log_audit_event(
            "pipeline_upload_profile",
            filename=file.filename or "dataset",
            size_bytes=len(body),
            dataset_file_id=data.get("dataset_file_id"),
            detail_level=detail_level,
        )
        return PipelineProfileStoredResponse.model_validate(data)
    except ValueError as exc:
        request_logger.error(
            "pipeline_upload_profile_failed filename=%r",
            file.filename or "dataset",
            exc_info=True,
        )
        raise HTTPException(status_code=400, detail={"code": "PIPELINE_PROFILE_ERROR", "message": str(exc)}) from exc


@app.post(f"{settings.api_prefix}/pipeline/profile-stored", response_model=PipelineProfileStoredResponse)
async def pipeline_profile_stored(payload: PipelineStoredProfileRequest) -> PipelineProfileStoredResponse:
    try:
        data = await fetch_stored_dataset_profile(
            dataset_file_id=payload.dataset_file_id,
            filename=payload.filename,
            histogram_bins=payload.histogram_bins,
            histogram_bins_by_field=payload.histogram_bins_by_field,
            detail_level=payload.profile_detail_level,
        )
        log_audit_event(
            "pipeline_profile_stored",
            dataset_file_id=payload.dataset_file_id,
            filename=payload.filename or "dataset",
            detail_level=payload.profile_detail_level,
        )
        return PipelineProfileStoredResponse.model_validate(data)
    except ValueError as exc:
        request_logger.error(
            "pipeline_profile_stored_failed dataset_file_id=%r filename=%r",
            payload.dataset_file_id,
            payload.filename or "dataset",
            exc_info=True,
        )
        raise HTTPException(status_code=400, detail={"code": "PIPELINE_PROFILE_ERROR", "message": str(exc)}) from exc


@app.post(f"{settings.api_prefix}/pipeline/run", response_model=PipelineRunResponse)
async def pipeline_run(
    file: UploadFile = File(...),
    config_json: str = Form(...),
    authorization: str | None = Header(default=None),
) -> PipelineRunResponse:
    try:
        config = PipelineConfig.model_validate_json(config_json)
        body = await file.read()
        request = PipelineRequest(filename=file.filename or "dataset", config=config)
        response = await run_pipeline_via_services(
            filename=file.filename or "dataset",
            body=body,
            payload=request,
            authorization=authorization,
        )
        log_audit_event(
            "pipeline_run",
            filename=file.filename or "dataset",
            size_bytes=len(body),
            mode=config.analysis_mode,
            target_row_id=config.target_row_id,
            criteria_count=len(config.criteria),
            history_id=response.history_id,
        )
        return response
    except ValueError as exc:
        request_logger.error(
            "pipeline_run_failed filename=%r",
            file.filename or "dataset",
            exc_info=True,
        )
        raise HTTPException(status_code=400, detail={"code": "PIPELINE_VALIDATION_ERROR", "message": str(exc)}) from exc


@app.post(f"{settings.api_prefix}/pipeline/run-stored", response_model=PipelineRunResponse)
async def pipeline_run_stored(
    payload: PipelineStoredRunRequest,
    authorization: str | None = Header(default=None),
) -> PipelineRunResponse:
    try:
        request = PipelineRequest(filename=payload.filename or "dataset", config=payload.config)
        response = await run_pipeline_from_storage(
            dataset_file_id=payload.dataset_file_id,
            filename=payload.filename,
            payload=request,
            authorization=authorization,
        )
        log_audit_event(
            "pipeline_run_stored",
            dataset_file_id=payload.dataset_file_id,
            filename=payload.filename or "dataset",
            mode=payload.config.analysis_mode,
            target_row_id=payload.config.target_row_id,
            criteria_count=len(payload.config.criteria),
            history_id=response.history_id,
        )
        return response
    except ValueError as exc:
        request_logger.error(
            "pipeline_run_stored_failed dataset_file_id=%r filename=%r",
            payload.dataset_file_id,
            payload.filename or "dataset",
            exc_info=True,
        )
        raise HTTPException(status_code=400, detail={"code": "PIPELINE_VALIDATION_ERROR", "message": str(exc)}) from exc


@app.post(f"{settings.api_prefix}/pipeline/preprocess-refresh", response_model=PipelinePreprocessRefreshResponse)
async def pipeline_preprocess_refresh(
    payload: PipelinePreprocessRefreshRequest,
    background_tasks: BackgroundTasks,
) -> PipelinePreprocessRefreshResponse:
    try:
        data = await refresh_preprocessing_from_storage(
            dataset_file_id=payload.dataset_file_id,
            filename=payload.filename,
            fields=[field.model_dump() for field in payload.fields],
            histogram_bins=payload.histogram_bins,
            histogram_bins_by_field=payload.histogram_bins_by_field,
            detail_level=payload.profile_detail_level,
        )
        if payload.profile_detail_level != "detailed":
            background_tasks.add_task(
                refresh_preprocessing_from_storage,
                dataset_file_id=payload.dataset_file_id,
                filename=payload.filename,
                fields=[field.model_dump() for field in payload.fields],
                histogram_bins=payload.histogram_bins,
                histogram_bins_by_field=payload.histogram_bins_by_field,
                detail_level="detailed",
            )
        log_audit_event(
            "pipeline_preprocess_refresh",
            dataset_file_id=payload.dataset_file_id,
            filename=payload.filename or "dataset",
            fields_count=len(payload.fields),
            histogram_bins=payload.histogram_bins,
            detail_level=payload.profile_detail_level,
        )
        return PipelinePreprocessRefreshResponse.model_validate(data)
    except ValueError as exc:
        request_logger.error(
            "pipeline_preprocess_refresh_failed dataset_file_id=%r filename=%r",
            payload.dataset_file_id,
            payload.filename or "dataset",
            exc_info=True,
        )
        raise HTTPException(
            status_code=400,
            detail={"code": "PIPELINE_PREPROCESSING_REFRESH_ERROR", "message": str(exc)},
        ) from exc


@app.post(f"{settings.api_prefix}/reports/comparison.docx")
def comparison_report(payload: ReportRequest) -> StreamingResponse:
    log_audit_event(
        "comparison_report_generated",
        title=payload.title,
        criteria_count=len(payload.criteria),
        ranking_count=len(payload.result.ranking),
    )
    document = Document()
    document.add_heading(payload.title, 0)
    document.add_paragraph("Отчет сформирован модульным монолитом системы сравнительного анализа.")

    document.add_heading("Сводка", level=1)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Показатель"
    summary_table.rows[0].cells[1].text = "Значение"
    for key, value in payload.result.analysis_summary.items():
        row = summary_table.add_row().cells
        row[0].text = report_label(str(key))
        row[1].text = report_value(str(key), value)

    document.add_heading("Критерии", level=1)
    criteria_table = document.add_table(rows=1, cols=4)
    criteria_table.style = "Table Grid"
    for cell, text_value in zip(criteria_table.rows[0].cells, ["Критерий", "Поле", "Вес", "Направление"]):
        cell.text = text_value
    for criterion in payload.criteria:
        row = criteria_table.add_row().cells
        row[0].text = criterion.name
        row[1].text = criterion.key
        row[2].text = f"{criterion.weight:.4f}"
        row[3].text = DIRECTION_LABELS.get(criterion.direction, criterion.direction)

    document.add_heading("Результаты", level=1)
    result_table = document.add_table(rows=1, cols=5)
    result_table.style = "Table Grid"
    for cell, text_value in zip(result_table.rows[0].cells, ["Место", "Объект", "Оценка", "Сходство", "Объяснение"]):
        cell.text = text_value
    for item in payload.result.ranking:
        row = result_table.add_row().cells
        row[0].text = str(item.rank)
        row[1].text = item.title
        row[2].text = f"{item.score:.4f}"
        row[3].text = "-" if item.similarity_to_target is None else f"{item.similarity_to_target:.4f}"
        row[4].text = item.explanation

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="comparison-report.docx"'},
    )


app.include_router(auth_router)
app.include_router(files_router, prefix=settings.api_prefix)
app.include_router(objects_router, prefix=settings.api_prefix)
