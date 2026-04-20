from io import BytesIO

import httpx
from fastapi import FastAPI, File, Form, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from docx import Document

from app.core_config import get_settings
from app.schemas.pipeline import (
    PipelineConfig,
    PipelinePreprocessRefreshRequest,
    PipelinePreprocessRefreshResponse,
    PipelineProfileStoredResponse,
    PipelineRequest,
    PipelineRunResponse,
    PipelineStoredRunRequest,
    ReportRequest,
)
from app.services.pipeline_engine import (
    fetch_system_dashboard,
    refresh_preprocessing_from_storage,
    run_pipeline_from_storage,
    run_pipeline_via_services,
    upload_and_profile_dataset,
)

settings = get_settings()

app = FastAPI(
    title=settings.app_name,
    version="0.1.0",
    debug=settings.debug,
    description="Единая точка входа, объединяющая импорт, предобработку и сравнительный анализ.",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_origin_list,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
def healthcheck() -> dict[str, str]:
    return {"status": "ok"}


SUMMARY_LABELS = {
    "objects_count": "Количество объектов",
    "criteria_count": "Количество критериев",
    "weights_sum": "Сумма весов",
    "best_object_id": "Лучший объект",
    "best_score": "Лучшая оценка",
    "normalization_notes": "Нормализация",
    "mode": "Режим анализа",
    "target_object_id": "Целевой объект",
    "confidence_score": "Доверие к расчету",
    "confidence_notes": "Замечания к качеству",
    "sensitivity": "Чувствительность",
    "ranking_stability_note": "Устойчивость рейтинга",
    "analog_groups": "Группы аналогов",
    "dominance_pairs": "Доминирование объектов",
}

DIRECTION_LABELS = {
    "maximize": "максимизация",
    "minimize": "минимизация",
    "target": "близость к целевому значению",
}

MODE_LABELS = {
    "rating": "рейтинг объектов",
    "analog_search": "поиск аналогов",
}


def report_label(key: str) -> str:
    return SUMMARY_LABELS.get(key, key.replace("_", " "))


def report_value(key: str, value: object) -> str:
    if value is None:
        return "-"
    if key == "mode":
        return MODE_LABELS.get(str(value), str(value))
    if isinstance(value, list):
        if not value:
            return "нет данных"
        if all(isinstance(item, str) for item in value):
            return "; ".join(str(item) for item in value)
        return f"{len(value)} записей"
    if isinstance(value, float):
        return f"{value:.4f}"
    if isinstance(value, dict):
        return "см. детальный раздел"
    return str(value)


@app.get("/api/v1/system/dashboard")
async def system_dashboard(authorization: str | None = Header(default=None)) -> dict:
    return await fetch_system_dashboard(settings=settings, authorization=authorization)


@app.post("/api/v1/pipeline/upload-profile", response_model=PipelineProfileStoredResponse)
async def pipeline_upload_profile(file: UploadFile = File(...)) -> PipelineProfileStoredResponse:
    try:
        body = await file.read()
        data = await upload_and_profile_dataset(settings=settings, filename=file.filename or "dataset", body=body)
        return PipelineProfileStoredResponse.model_validate(data)
    except (ValueError, httpx.HTTPError) as exc:
        raise HTTPException(
            status_code=400,
            detail={"code": "PIPELINE_PROFILE_ERROR", "message": str(exc)},
        ) from exc


@app.post("/api/v1/pipeline/run", response_model=PipelineRunResponse)
async def pipeline_run(
    file: UploadFile = File(...),
    config_json: str = Form(...),
    authorization: str | None = Header(default=None),
) -> PipelineRunResponse:
    try:
        config = PipelineConfig.model_validate_json(config_json)
        body = await file.read()
        request = PipelineRequest(filename=file.filename or "dataset", config=config)
        return await run_pipeline_via_services(
            settings=settings,
            filename=file.filename or "dataset",
            body=body,
            payload=request,
            authorization=authorization,
        )
    except (ValueError, httpx.HTTPError) as exc:
        raise HTTPException(
            status_code=400,
            detail={"code": "PIPELINE_VALIDATION_ERROR", "message": str(exc)},
        ) from exc


@app.post("/api/v1/pipeline/run-stored", response_model=PipelineRunResponse)
async def pipeline_run_stored(
    payload: PipelineStoredRunRequest,
    authorization: str | None = Header(default=None),
) -> PipelineRunResponse:
    try:
        request = PipelineRequest(filename=payload.filename or "dataset", config=payload.config)
        return await run_pipeline_from_storage(
            settings=settings,
            dataset_file_id=payload.dataset_file_id,
            filename=payload.filename,
            payload=request,
            authorization=authorization,
        )
    except (ValueError, httpx.HTTPError) as exc:
        raise HTTPException(
            status_code=400,
            detail={"code": "PIPELINE_VALIDATION_ERROR", "message": str(exc)},
        ) from exc


@app.post("/api/v1/pipeline/preprocess-refresh", response_model=PipelinePreprocessRefreshResponse)
async def pipeline_preprocess_refresh(payload: PipelinePreprocessRefreshRequest) -> PipelinePreprocessRefreshResponse:
    try:
        data = await refresh_preprocessing_from_storage(
            settings=settings,
            dataset_file_id=payload.dataset_file_id,
            filename=payload.filename,
            fields=[field.model_dump() for field in payload.fields],
            histogram_bins=payload.histogram_bins,
            histogram_bins_by_field=payload.histogram_bins_by_field,
        )
        return PipelinePreprocessRefreshResponse.model_validate(data)
    except (ValueError, httpx.HTTPError) as exc:
        raise HTTPException(
            status_code=400,
            detail={"code": "PIPELINE_PREPROCESSING_REFRESH_ERROR", "message": str(exc)},
        ) from exc


@app.post("/api/v1/reports/comparison.docx")
def comparison_report(payload: ReportRequest) -> StreamingResponse:
    document = Document()
    document.add_heading(payload.title, 0)
    document.add_paragraph("Сформировано информационно-аналитической системой сравнительного анализа объектов.")
    document.add_heading("Сводка", level=1)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Показатель"
    summary_table.rows[0].cells[1].text = "Значение"
    for key, value in payload.result.analysis_summary.items():
        row = summary_table.add_row().cells
        row[0].text = report_label(str(key))
        row[1].text = report_value(str(key), value)

    document.add_heading("Критерии", level=1)
    criteria_table = document.add_table(rows=1, cols=4)
    criteria_table.style = "Table Grid"
    for cell, text in zip(criteria_table.rows[0].cells, ["Критерий", "Поле", "Вес", "Направление"]):
        cell.text = text
    for criterion in payload.criteria:
        row = criteria_table.add_row().cells
        row[0].text = criterion.name
        row[1].text = criterion.key
        row[2].text = f"{criterion.weight:.4f}"
        row[3].text = DIRECTION_LABELS.get(criterion.direction, criterion.direction)

    document.add_heading("Результаты", level=1)
    result_table = document.add_table(rows=1, cols=5)
    result_table.style = "Table Grid"
    for cell, text in zip(result_table.rows[0].cells, ["Место", "Объект", "Оценка", "Сходство", "Объяснение"]):
        cell.text = text
    for item in payload.result.ranking:
        row = result_table.add_row().cells
        row[0].text = str(item.rank)
        row[1].text = item.title
        row[2].text = f"{item.score:.4f}"
        row[3].text = "-" if item.similarity_to_target is None else f"{item.similarity_to_target:.4f}"
        row[4].text = item.explanation

    document.add_heading("Вклад критериев для лучшего результата", level=1)
    if payload.result.ranking:
        contribution_table = document.add_table(rows=1, cols=5)
        contribution_table.style = "Table Grid"
        for cell, text in zip(contribution_table.rows[0].cells, ["Критерий", "Исходное значение", "Норм.", "Вес", "Вклад"]):
            cell.text = text
        for contribution in payload.result.ranking[0].contributions:
            row = contribution_table.add_row().cells
            row[0].text = contribution.name
            row[1].text = str(contribution.raw_value)
            row[2].text = f"{contribution.normalized_value:.4f}"
            row[3].text = f"{contribution.weight:.4f}"
            row[4].text = f"{contribution.contribution:.4f}"

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="comparison-report.docx"'},
    )
